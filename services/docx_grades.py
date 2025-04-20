from docx import Document
from io import BytesIO
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def generate_grades_certificate(data: dict, template_path: str = "templates/grades_template.docx") -> BytesIO:
    doc = Document(template_path)

    # Set current date if not already included
    if "date_today" not in data:
        data["date_today"] = datetime.now().strftime("%B %d, %Y")  # e.g., April 19, 2025

    # Replace placeholders in full paragraph text
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        for key, value in data.items():
            if key != "subjects":
                placeholder = f"{{{{{key}}}}}"
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))

        if paragraph.text != full_text:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)
    # Subjects Table Handling
    if "subjects" in data:
        subjects = data["subjects"]
        table = doc.tables[-1]
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[1]._tr)

        sum_of_grades = 0
        number_of_subjects = len(subjects)

        for subject in subjects:
            grade = subject["final_grade"]
            units = subject["units"]

            sum_of_grades += grade

            row = table.add_row()
            row.cells[0].text = subject["subject_code"]
            row.cells[1].text = subject["subject_name"]
            row.cells[2].text = f"{grade:.1f}"
            row.cells[3].text = str(units)

            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Calculate and append GWA
        gwa = sum_of_grades / number_of_subjects if number_of_subjects else 0.0
        gwa_row = table.add_row()
        gwa_cell_2 = gwa_row.cells[2]
        gwa_paragraph_2 = gwa_cell_2.paragraphs[0]
        gwa_run_2 = gwa_paragraph_2.add_run("GWA")
        gwa_run_2.bold = True

        # Cell 3 - GWA value
        gwa_cell_3 = gwa_row.cells[3]
        gwa_paragraph_3 = gwa_cell_3.paragraphs[0]
        gwa_run_3 = gwa_paragraph_3.add_run(f"{gwa:.1f}")
        gwa_run_3.bold = True
        

        for cell in gwa_row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
